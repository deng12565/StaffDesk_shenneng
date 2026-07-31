"""StaffDeck 后端模块：知识文档文本提取器，支持文本、HTML、PDF 和 DOCX 等格式。

主要入口：KnowledgeParseError, extract_text。阅读时先从这些入口跟踪调用关系。
"""

from __future__ import annotations

import re
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile


SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".html", ".htm", ".pdf", ".docx", ".doc"}


class KnowledgeParseError(ValueError):
    pass


def extract_text(filename: str, content: bytes) -> tuple[str, str]:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise KnowledgeParseError(f"暂不支持 {suffix or 'unknown'} 文件格式。")
    if suffix == ".doc":
        raise KnowledgeParseError("暂不支持旧版 .doc 二进制格式，请转换为 .docx 后上传。")
    if suffix in {".txt", ".md", ".markdown"}:
        return _decode_text(content), suffix.lstrip(".")
    if suffix in {".html", ".htm"}:
        return _extract_html(content), "html"
    if suffix == ".pdf":
        return _extract_pdf(content), "pdf"
    if suffix == ".docx":
        return _extract_docx(content), "docx"
    raise KnowledgeParseError(f"暂不支持 {suffix} 文件格式。")


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _extract_html(content: bytes) -> str:
    text = _decode_text(content)
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(text, "html.parser")
        for item in soup(["script", "style", "noscript"]):
            item.decompose()
        return soup.get_text("\n")
    except Exception:
        parser = _HTMLTextExtractor()
        parser.feed(text)
        return parser.text


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - dependency availability differs by env.
        raise KnowledgeParseError("缺少 pypdf，无法解析 PDF。") from exc
    reader = PdfReader(BytesIO(content))
    pages: list[str] = []
    for index, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"## 第 {index + 1} 页\n\n{page_text}")
    if not pages:
        return ""
    return "# PDF 文档\n\n" + "\n\n".join(pages)


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document

        document = Document(BytesIO(content))
        rows: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            heading_level = _docx_heading_level(paragraph.style.name if paragraph.style else "")
            rows.append(f"{'#' * heading_level} {text}" if heading_level else text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
        return "\n".join(rows)
    except Exception:
        return _extract_docx_with_zip(content)


def _docx_heading_level(style_name: str) -> int | None:
    match = re.match(r"^(?:Heading|标题)\s*([1-6])$", style_name.strip(), re.IGNORECASE)
    return int(match.group(1)) if match else None


def _extract_docx_with_zip(content: bytes) -> str:
    try:
        with ZipFile(BytesIO(content)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    except Exception as exc:
        raise KnowledgeParseError("无法解析 docx 文档。") from exc
    parser = _DocxTextExtractor()
    parser.feed(xml)
    return parser.text


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []

    @property
    def text(self) -> str:
        return "\n".join(part.strip() for part in self._parts if part.strip())

    def handle_data(self, data: str) -> None:
        if data.strip():
            self._parts.append(data)


class _DocxTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []

    @property
    def text(self) -> str:
        return "\n".join(part.strip() for part in self._parts if part.strip())

    def handle_data(self, data: str) -> None:
        if data.strip():
            self._parts.append(data)
