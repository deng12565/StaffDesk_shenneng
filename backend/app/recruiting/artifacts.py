from __future__ import annotations

import base64
from functools import lru_cache
import hashlib
import io
import json
import os
from pathlib import Path, PurePosixPath, PureWindowsPath
import re
import subprocess
import sys
import tempfile
from typing import Any
import zipfile

from cryptography.fernet import Fernet
from docx import Document
from pypdf import PdfReader

from app import paths
from app.config import get_settings


ALLOWED_FORMATS = {"pdf", "docx", "doc", "docm", "jpeg", "png", "webp", "zip", "rar", "7z"}
ARCHIVE_FORMATS = {"zip", "rar", "7z"}
MAX_ARCHIVE_ENTRIES = 100
MAX_ARCHIVE_EXPANDED_BYTES = 100 * 1024 * 1024
MAX_ARCHIVE_RATIO = 100
MAX_ARCHIVE_DEPTH = 2


class ArtifactError(RuntimeError):
    def __init__(self, code: str, message: str = ""):
        super().__init__(message or code)
        self.code = code


class ArtifactStore:
    def __init__(self, root: Path | None = None):
        settings = get_settings()
        configured = str(settings.recruiting_data_dir or "").strip()
        self.root = root or (
            Path(configured).expanduser().resolve()
            if configured
            else (paths.user_data_dir() / "recruiting").resolve()
        )

    def put(self, tenant_id: str, application_id: str, payload: bytes) -> str:
        digest = hashlib.sha256(payload).hexdigest()
        folder = self.root / _safe_segment(tenant_id) / _safe_segment(application_id)
        folder.mkdir(parents=True, exist_ok=True)
        target = folder / f"{digest}.bin.enc"
        encrypted = Fernet(_artifact_key()).encrypt(payload)
        temporary = target.with_suffix(".tmp")
        temporary.write_bytes(encrypted)
        os.replace(temporary, target)
        return str(target.relative_to(self.root)).replace("\\", "/")

    def get(self, reference: str) -> bytes:
        target = (self.root / reference).resolve()
        if self.root not in target.parents:
            raise ArtifactError("ATTACHMENT_QUARANTINED")
        return Fernet(_artifact_key()).decrypt(target.read_bytes())

    def delete(self, reference: str) -> bool:
        target = (self.root / reference).resolve()
        if self.root not in target.parents or not target.exists():
            return False
        target.unlink()
        return True


def detect_format(filename: str, content_type: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower().lstrip(".")
    signatures = {
        "pdf": data.startswith(b"%PDF-"),
        "doc": data.startswith(bytes.fromhex("D0CF11E0A1B11AE1")),
        "zip": data.startswith(b"PK\x03\x04") or data.startswith(b"PK\x05\x06"),
        "rar": data.startswith(b"Rar!\x1a\x07"),
        "7z": data.startswith(b"7z\xbc\xaf\x27\x1c"),
        "jpeg": data.startswith(b"\xff\xd8\xff"),
        "png": data.startswith(b"\x89PNG\r\n\x1a\n"),
        "webp": len(data) >= 12 and data[:4] == b"RIFF" and data[8:12] == b"WEBP",
    }
    if signatures["zip"] and suffix in {"docx", "docm"}:
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as package:
                if "[Content_Types].xml" in package.namelist() and any(
                    name.startswith("word/") for name in package.namelist()
                ):
                    return suffix
        except zipfile.BadZipFile as exc:
            raise ArtifactError("PARSE_FAILED") from exc
    detected = next((name for name, matched in signatures.items() if matched), "")
    expected = "rar" if suffix in {"rar", "rar5"} else suffix
    if expected not in ALLOWED_FORMATS:
        raise ArtifactError("ATTACHMENT_UNSUPPORTED")
    if not detected or detected != expected:
        raise ArtifactError("ATTACHMENT_QUARANTINED", f"signature mismatch for {content_type}")
    return detected


def extract_text(file_format: str, data: bytes, *, filename: str = "attachment") -> str:
    if file_format == "pdf":
        try:
            reader = PdfReader(io.BytesIO(data))
            if reader.is_encrypted:
                raise ArtifactError("ATTACHMENT_UNSUPPORTED", "encrypted PDF")
            if len(reader.pages) > 30:
                raise ArtifactError("ARCHIVE_LIMIT_EXCEEDED", "PDF exceeds 30 pages")
            text = "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
        except ArtifactError:
            raise
        except Exception as exc:
            raise ArtifactError("PARSE_FAILED") from exc
        if not text:
            raise ArtifactError("VISION_UNAVAILABLE", "scanned PDF requires a verified vision model")
        return text
    if file_format == "docx":
        try:
            document = Document(io.BytesIO(data))
            parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text.strip() for cell in row.cells))
            return "\n".join(part for part in parts if part.strip()).strip()
        except Exception as exc:
            raise ArtifactError("PARSE_FAILED") from exc
    if file_format in {"jpeg", "png", "webp"}:
        try:
            from PIL import Image

            with Image.open(io.BytesIO(data)) as image:
                image.verify()
        except Exception as exc:
            raise ArtifactError("PARSE_FAILED") from exc
        raise ArtifactError("VISION_UNAVAILABLE", "image resume requires a verified vision model")
    if file_format in {"doc", "docm"}:
        converted = convert_word_document(data, filename)
        return extract_text("docx", converted, filename=f"{filename}.docx")
    raise ArtifactError("ATTACHMENT_UNSUPPORTED")


def convert_word_document(data: bytes, filename: str) -> bytes:
    with tempfile.TemporaryDirectory(prefix="staffdeck-recruiting-word-") as folder:
        root = Path(folder)
        source = root / (Path(filename).name or "resume.doc")
        output = root / "converted.docx"
        source.write_bytes(data)
        try:
            completed = subprocess.run(
                [
                    sys.executable,
                    "-m",
                    "app.recruiting.word_converter",
                    str(source),
                    str(output),
                ],
                stdin=subprocess.DEVNULL,
                capture_output=True,
                text=True,
                timeout=60,
                check=False,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            raise ArtifactError("DOCUMENT_CONVERTER_UNAVAILABLE") from exc
        if completed.returncode != 0 or not output.exists() or output.stat().st_size == 0:
            raise ArtifactError("DOCUMENT_CONVERSION_FAILED")
        payload = output.read_bytes()
    _ensure_macro_free_docx(payload)
    return payload


def list_archive_entries(archive_path: Path, *, seven_zip_path: str | None = None) -> list[dict[str, Any]]:
    executable = seven_zip_path or get_settings().recruiting_7z_path
    try:
        completed = subprocess.run(
            [executable, "l", "-slt", "-ba", "-sccUTF-8", str(archive_path)],
            stdin=subprocess.DEVNULL,
            capture_output=True,
            timeout=60,
            check=False,
        )
    except OSError as exc:
        raise ArtifactError("ARCHIVE_EXTRACTOR_UNAVAILABLE") from exc
    text = completed.stdout.decode("utf-8", errors="replace")
    if completed.returncode != 0:
        if "password" in text.lower() or "encrypted" in text.lower():
            raise ArtifactError("ARCHIVE_ENCRYPTED")
        raise ArtifactError("ARCHIVE_EXTRACTION_FAILED")
    entries = _parse_7z_slt(text)
    _validate_archive_entries(entries)
    return entries


def extract_archive_entries(
    data: bytes,
    filename: str,
    *,
    depth: int = 0,
    seven_zip_path: str | None = None,
) -> list[tuple[str, bytes]]:
    if depth >= MAX_ARCHIVE_DEPTH:
        raise ArtifactError("ARCHIVE_LIMIT_EXCEEDED")
    executable = seven_zip_path or get_settings().recruiting_7z_path
    suffix = Path(filename).suffix or ".bin"
    with tempfile.TemporaryDirectory(prefix="staffdeck-recruiting-archive-") as folder:
        archive_path = Path(folder) / f"archive{suffix}"
        archive_path.write_bytes(data)
        entries = list_archive_entries(archive_path, seven_zip_path=executable)
        test = subprocess.run(
            [executable, "t", "-y", "-bd", str(archive_path)],
            stdin=subprocess.DEVNULL,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=60,
            check=False,
        )
        if test.returncode != 0:
            raise ArtifactError("ARCHIVE_EXTRACTION_FAILED")
        extracted: list[tuple[str, bytes]] = []
        total = 0
        for entry in entries:
            path = str(entry["Path"])
            if entry.get("Folder") == "+":
                continue
            file_format = Path(path).suffix.lower().lstrip(".")
            file_format = "rar" if file_format == "rar5" else file_format
            if file_format not in ALLOWED_FORMATS:
                raise ArtifactError("ARCHIVE_UNSAFE_ENTRY")
            try:
                item = subprocess.run(
                    [executable, "x", "-so", "-y", "-bd", str(archive_path), path],
                    stdin=subprocess.DEVNULL,
                    capture_output=True,
                    timeout=60,
                    check=False,
                )
            except subprocess.TimeoutExpired as exc:
                raise ArtifactError("ARCHIVE_LIMIT_EXCEEDED") from exc
            if item.returncode != 0:
                raise ArtifactError("ARCHIVE_EXTRACTION_FAILED")
            if len(item.stdout) > get_settings().recruiting_max_attachment_bytes:
                raise ArtifactError("ARCHIVE_LIMIT_EXCEEDED")
            total += len(item.stdout)
            if total > MAX_ARCHIVE_EXPANDED_BYTES:
                raise ArtifactError("ARCHIVE_LIMIT_EXCEEDED")
            extracted.append((path, item.stdout))
        return extracted


@lru_cache(maxsize=1)
def probe_document_capabilities() -> dict[str, Any]:
    result: dict[str, Any] = {
        "word": {"available": False, "version": None, "error_code": None},
        "seven_zip": {"available": False, "version": None, "formats": [], "error_code": None},
    }
    try:
        probe = subprocess.run(
            [sys.executable, "-m", "app.recruiting.word_converter", "--probe"],
            stdin=subprocess.DEVNULL,
            capture_output=True,
            text=True,
            timeout=20,
            check=False,
        )
        if probe.returncode == 0:
            result["word"].update(json.loads(probe.stdout))
        else:
            result["word"]["error_code"] = "DOCUMENT_CONVERTER_UNAVAILABLE"
    except (OSError, subprocess.TimeoutExpired, json.JSONDecodeError):
        result["word"]["error_code"] = "DOCUMENT_CONVERTER_UNAVAILABLE"
    executable = get_settings().recruiting_7z_path
    try:
        probe = subprocess.run(
            [executable, "i"],
            stdin=subprocess.DEVNULL,
            capture_output=True,
            timeout=20,
            check=False,
        )
        text = probe.stdout.decode("utf-8", errors="replace")
        version = next((line.strip() for line in text.splitlines() if "7-Zip" in line), None)
        formats = [name for name in ("zip", "rar", "7z") if re.search(rf"\b{name}\b", text, re.I)]
        available = probe.returncode == 0 and set(formats) == {"zip", "rar", "7z"}
        result["seven_zip"].update({"available": available, "version": version, "formats": formats})
        if not available:
            result["seven_zip"]["error_code"] = "ARCHIVE_EXTRACTOR_UNAVAILABLE"
    except (OSError, subprocess.TimeoutExpired):
        result["seven_zip"]["error_code"] = "ARCHIVE_EXTRACTOR_UNAVAILABLE"
    return result


def _parse_7z_slt(text: str) -> list[dict[str, str]]:
    entries: list[dict[str, str]] = []
    current: dict[str, str] = {}
    for line in text.splitlines():
        if not line.strip():
            if current.get("Path"):
                entries.append(current)
            current = {}
            continue
        key, separator, value = line.partition(" = ")
        if separator:
            current[key.strip()] = value.strip()
    if current.get("Path"):
        entries.append(current)
    return entries


def _validate_archive_entries(entries: list[dict[str, str]]) -> None:
    if len(entries) > MAX_ARCHIVE_ENTRIES:
        raise ArtifactError("ARCHIVE_LIMIT_EXCEEDED")
    total_size = 0
    total_packed = 0
    for entry in entries:
        name = str(entry.get("Path") or "")
        posix = PurePosixPath(name.replace("\\", "/"))
        windows = PureWindowsPath(name)
        attributes = str(entry.get("Attributes") or "").upper()
        if (
            not name
            or posix.is_absolute()
            or windows.is_absolute()
            or ".." in posix.parts
            or any(":" in part for part in posix.parts)
            or "L" in attributes
            or entry.get("Symbolic Link")
            or entry.get("Hard Link")
        ):
            raise ArtifactError("ARCHIVE_UNSAFE_ENTRY")
        if str(entry.get("Encrypted") or "-") == "+":
            raise ArtifactError("ARCHIVE_ENCRYPTED")
        size = int(entry.get("Size") or 0)
        packed = int(entry.get("Packed Size") or 0)
        if size > get_settings().recruiting_max_attachment_bytes:
            raise ArtifactError("ARCHIVE_LIMIT_EXCEEDED")
        if packed > 0 and size / packed > MAX_ARCHIVE_RATIO:
            raise ArtifactError("ARCHIVE_LIMIT_EXCEEDED")
        total_size += size
        total_packed += packed
    if total_size > MAX_ARCHIVE_EXPANDED_BYTES:
        raise ArtifactError("ARCHIVE_LIMIT_EXCEEDED")
    if total_packed > 0 and total_size / total_packed > MAX_ARCHIVE_RATIO:
        raise ArtifactError("ARCHIVE_LIMIT_EXCEEDED")


def _ensure_macro_free_docx(payload: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as package:
            names = {name.lower() for name in package.namelist()}
            content_types = package.read("[Content_Types].xml").lower()
    except (KeyError, zipfile.BadZipFile) as exc:
        raise ArtifactError("DOCUMENT_CONVERSION_FAILED") from exc
    if any("vbaproject" in name for name in names) or b"macroenabled" in content_types:
        raise ArtifactError("DOCUMENT_CONVERSION_FAILED")


def _artifact_key() -> bytes:
    material = f"{get_settings().app_secret}:recruiting-artifacts".encode("utf-8")
    return base64.urlsafe_b64encode(hashlib.sha256(material).digest())


def _safe_segment(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]", "_", value)[:100]
